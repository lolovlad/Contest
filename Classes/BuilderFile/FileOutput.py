import openpyxl
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from os import remove
from .Cell import Cell


class FileXlsx:
    def __init__(self, workbook: openpyxl.Workbook = openpyxl.Workbook()):
        self.__cells_name: dict = {}
        self.__main_body_cells = []
        self.__workbook = workbook
        self.__sheet = self.__workbook.active

    def add_main_body(self, cell: Cell):
        self.__main_body_cells.append(cell)

    def add_cells_name(self, cell: Cell):
        self.__cells_name[cell.value] = cell

    def return_file(self, name_file: str):
        for cell in list(self.__cells_name.values()) + self.__main_body_cells:
            cell.create_cell(self.__sheet)
        self.__workbook.save(f"{name_file}.xlsx")

    def find_to_name_cell(self, name_cell: str):
        return self.__cells_name[name_cell]


class FileDocx:
    def __init__(self, workbook: Document = Document()):
        self.__cells_name: dict = {}
        self.__main_body_cells = []
        self.__workbook: Document = workbook

    def add_main_body(self, cell: Cell):
        self.__main_body_cells.append(cell)

    def add_cells_name(self, cell: Cell):
        self.__cells_name[cell.value] = cell

    def return_file(self, name_file: str):
        last_col = list(self.__cells_name.values())[-1].col
        last_row = self.__main_body_cells[-1].row

        table = self.__workbook.add_table(rows=last_row, cols=last_col, style="Table Grid")
        self.settings_tabel()
        for cell in list(self.__cells_name.values()):
            cell.create_cell(table)
        for cell in self.__main_body_cells:
            cell.create_cell(table)
        self.__workbook.save(f"{name_file}.docx")

    def settings_tabel(self):
        pass

    def find_to_name_cell(self, name_cell: str):
        return self.__cells_name[name_cell]


class FilePDF:
    def __init__(self, workbook: Document = Document()):
        self.__cells_name: dict = {}
        self.__main_body_cells = []
        self.__workbook: Document = workbook

    def add_main_body(self, cell: Cell):
        self.__main_body_cells.append(cell)

    def add_cells_name(self, cell: Cell):
        self.__cells_name[cell.value] = cell

    def return_file(self, name_file: str):
        last_col = list(self.__cells_name.values())[-1].col
        last_row = self.__main_body_cells[-1].row

        table = self.__workbook.add_table(rows=last_row, cols=last_col, style="Table Grid")
        self.settings_tabel()
        for cell in list(self.__cells_name.values()):
            cell.create_cell(table)
        for cell in self.__main_body_cells:
            cell.create_cell(table)
        self.__workbook.save(f"convert_docx_file.docx")
        convert(f"convert_docx_file.docx", f"{name_file}.pdf")
        remove(f"convert_docx_file.docx")

    def settings_tabel(self):
        pass

    def find_to_name_cell(self, name_cell: str):
        return self.__cells_name[name_cell]